import os
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Set the background color of a cell (fill_color as hex string like 'F2F2F2')."""
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def add_page_number(run):
    """Add page number field to a footer run."""
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

def add_toc(paragraph):
    """Add a dynamic Table of Contents field to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    placeholder = paragraph.add_run("Right-click this field and select 'Update Field' to generate Table of Contents.")
    placeholder.font.italic = True
    placeholder.font.color.rgb = RGBColor(128, 128, 128)
    
    run2 = paragraph.add_run()
    run2._r.append(fldChar3)

def add_horizontal_rule(doc):
    """Add a thin horizontal divider line using a 1-row, 1-col table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "CCCCCC")
    
    # Set height to 1pt (20 dxa)
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trHeight = parse_xml(r'<w:trHeight %s w:val="20" w:hRule="exact"/>' % nsdecls('w'))
    trPr.append(trHeight)
    
    # Remove borders
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    # Minimize cell padding
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="0" w:type="dxa"/>'
        f'<w:left w:w="0" w:type="dxa"/>'
        f'<w:bottom w:w="0" w:type="dxa"/>'
        f'<w:right w:w="0" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)
    
    # Clear spacing on cell's paragraph
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def parse_and_add_runs(paragraph, text):
    """Parse basic markdown inline styles (**bold**, *italic*, `code`) and append runs."""
    # Match: **bold**, *italic*, _italic_, `code`
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|_.*?_)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("_") and part.endswith("_"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(197, 17, 98) # Pinkish-red standard inline code style
        else:
            run = paragraph.add_run(part)

def add_code_block(doc, code_text, lang=""):
    """Render a code block in a beautiful single-cell table callout."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    # Set background shading to very light gray
    set_cell_background(cell, "F8F9FA")
    
    # Set margins for code block cell (left margin indented, top/bottom padding)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="120" w:type="dxa"/>'
        f'<w:left w:w="180" w:type="dxa"/>'
        f'<w:bottom w:w="120" w:type="dxa"/>'
        f'<w:right w:w="180" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)
    
    # Set a light grey thin border
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="E4E7EC"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="A4BCDE"/>' # Thicker blue-gray left border
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E4E7EC"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="E4E7EC"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(34, 34, 34)
    
    # Spacing after code block
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

def add_markdown_table(doc, table_lines):
    """Parse and add a markdown table with headers and styled grid."""
    rows_data = []
    for line in table_lines:
        line_strip = line.strip()
        if not line_strip:
            continue
        cells = [c.strip() for c in line_strip.split("|")]
        if len(cells) >= 2:
            if line_strip.startswith("|"):
                cells = cells[1:]
            if line_strip.endswith("|"):
                cells = cells[:-1]
            rows_data.append(cells)
            
    if not rows_data:
        return
        
    # Check if second row is separator (e.g. |---|---|)
    if len(rows_data) > 1 and all(re.match(r'^:?-+:?$', c) for c in rows_data[1]):
        header = rows_data[0]
        rows = rows_data[2:]
    else:
        header = None
        rows = rows_data
        
    num_cols = max(len(r) for r in ([header] if header else []) + rows)
    num_rows = (1 if header else 0) + len(rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    row_idx = 0
    if header:
        for col_idx, text in enumerate(header):
            if col_idx < len(table.rows[row_idx].cells):
                cell = table.cell(row_idx, col_idx)
                set_cell_background(cell, "1F4E79") # Navy Primary
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        row_idx += 1
        
    for r_data in rows:
        for col_idx, text in enumerate(r_data):
            if col_idx < len(table.rows[row_idx].cells):
                cell = table.cell(row_idx, col_idx)
                if row_idx % 2 == 1:
                    set_cell_background(cell, "F2F5F8") # Zebra striping
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                parse_and_add_runs(p, text)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        row_idx += 1
        
    # Spacer after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)

def add_image(doc, img_path, alt_text):
    """Load and center an image, adding a structured caption underneath."""
    base_dir = r"a:\Interview_Generator"
    full_path = os.path.join(base_dir, img_path.replace("/", "\\"))
    
    if os.path.exists(full_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(full_path, width=Inches(5.8)) # Scale to fit margins cleanly
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run(f"Figure: {alt_text}")
        run_cap.italic = True
        run_cap.font.size = Pt(9.5)
        run_cap.font.color.rgb = RGBColor(100, 100, 100)
    else:
        print(f"Warning: Image not found at {full_path}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Image Missing: {alt_text} at {img_path}]")
        run.bold = True
        run.font.color.rgb = RGBColor(200, 0, 0)

def add_cover_page(doc):
    """Add a professional academic project report cover page."""
    # Vertical spacing
    for _ in range(2):
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
    p_rep = doc.add_paragraph()
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_rep = p_rep.add_run("PROJECT REPORT ON")
    run_rep.font.name = "Calibri"
    run_rep.font.size = Pt(12)
    run_rep.bold = True
    run_rep.font.color.rgb = RGBColor(80, 80, 80)
    p_rep.paragraph_format.space_after = Pt(12)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PREPEDGE AI: AN AI-POWERED INTERVIEW SIMULATOR &\nMOCK EVALUATION PLATFORM")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(22)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121) # Navy Primary
    p_title.paragraph_format.space_after = Pt(18)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(
        "Submitted in partial fulfillment of the requirements\n"
        "for the degree of\n"
        "Bachelor of Computer Applications (BCA)"
    )
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(11)
    run_sub.bold = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    p_sub.paragraph_format.space_after = Pt(36)
    
    add_horizontal_rule(doc)
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(36)
    
    # Metadata grid table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    for row in table.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(3.0)
            
    meta = [
        ("Submitted By:", "[Your Name]\nRoll No: [Your Roll Number]"),
        ("Guided By:", "[Supervisor Name]\nAssistant Professor"),
        ("Department:", "Department of Computer Applications"),
        ("Institution:", "[Your College/Institution Name]\nAcademic Year: 2025-2026")
    ]
    
    for idx, (label, val) in enumerate(meta):
        cell_lbl = table.cell(idx, 0)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(4)
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.name = "Calibri"
        run_lbl.font.size = Pt(10.5)
        run_lbl.bold = True
        run_lbl.font.color.rgb = RGBColor(80, 80, 80)
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        cell_val = table.cell(idx, 1)
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(4)
        run_val = p_val.add_run(val)
        run_val.font.name = "Calibri"
        run_val.font.size = Pt(10.5)
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="none"/>'
                f'<w:left w:val="none"/>'
                f'<w:bottom w:val="none"/>'
                f'<w:right w:val="none"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)
            
    doc.add_page_break()

def configure_styles(doc):
    """Set global document styles for Calibri consistency."""
    styles = doc.styles
    
    # Normal style
    style_normal = styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(34, 34, 34)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Heading 1
    style_h1 = styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Calibri'
    font_h1.size = Pt(18)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(31, 78, 121) # Navy Primary
    style_h1.paragraph_format.space_before = Pt(18)
    style_h1.paragraph_format.space_after = Pt(8)
    style_h1.paragraph_format.keep_with_next = True
    
    # Heading 2
    style_h2 = styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Calibri'
    font_h2.size = Pt(14)
    font_h2.bold = True
    font_h2.color.rgb = RGBColor(47, 85, 151) # Navy Secondary
    style_h2.paragraph_format.space_before = Pt(14)
    style_h2.paragraph_format.space_after = Pt(6)
    style_h2.paragraph_format.keep_with_next = True
    
    # Heading 3
    style_h3 = styles['Heading 3']
    font_h3 = style_h3.font
    font_h3.name = 'Calibri'
    font_h3.size = Pt(12)
    font_h3.bold = True
    font_h3.color.rgb = RGBColor(89, 89, 89) # Slate Gray
    style_h3.paragraph_format.space_before = Pt(10)
    style_h3.paragraph_format.space_after = Pt(4)
    style_h3.paragraph_format.keep_with_next = True

def main():
    md_path = r"a:\Interview_Generator\PrepEdge_AI_Project_Report.md"
    docx_path = r"a:\Interview_Generator\PrepEdge_AI_Project_Report.docx"
    
    print(f"Reading markdown from: {md_path}")
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    lines = content.splitlines()
    
    print("Initializing docx document...")
    doc = Document()
    configure_styles(doc)
    
    # Configure document geometry (Margins: 1 inch on all sides)
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Enable different cover page header/footer
    section.different_first_page_header_footer = True
    
    # Setup standard header/footer
    header = section.header
    p_header = header.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = p_header.add_run("PrepEdge AI — Project Report")
    run_h.font.name = "Calibri"
    run_h.font.size = Pt(8.5)
    run_h.font.color.rgb = RGBColor(128, 128, 128)
    
    footer = section.footer
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run()
    add_page_number(run_f)
    run_f.font.name = "Calibri"
    run_f.font.size = Pt(9.5)
    run_f.font.color.rgb = RGBColor(128, 128, 128)
    
    print("Generating Cover Page...")
    add_cover_page(doc)
    
    print("Parsing main markdown body...")
    i = 0
    in_front_matter = True
    
    # We skipped the first 7 lines in raw parsing because they are represented on the cover page.
    # Let's clean up index matching to start from Line 8 (index 7).
    i = 7
    
    while i < len(lines):
        line = lines[i]
        
        # 1. Skip Table of Contents block to insert a dynamic one
        if line.strip() == "## TABLE OF CONTENTS":
            i += 1
            # Skip everything until next "---"
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            # Insert a clean dynamic Table of Contents page
            p_toc_head = doc.add_heading("TABLE OF CONTENTS", level=1)
            p_toc_head.paragraph_format.page_break_before = False # Will rely on natural structure
            p_toc_head.paragraph_format.space_before = Pt(12)
            p_toc_head.paragraph_format.space_after = Pt(12)
            
            p_toc = doc.add_paragraph()
            add_toc(p_toc)
            doc.add_page_break()
            
            # Skip the trailing "---" line too
            if i < len(lines) and lines[i].strip() == "---":
                i += 1
            continue
            
        # 2. Check if we have transitioned past the front matter
        if "## 1. INTRODUCTION & OBJECTIVES" in line:
            in_front_matter = False
            
        # 3. Code block parsing
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            
            print(f"Adding code block: {lang} (length: {len(code_lines)} lines)")
            add_code_block(doc, "\n".join(code_lines), lang)
            i += 1
            continue
            
        # 4. Table parsing
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            print(f"Adding table ({len(table_lines)} rows)")
            add_markdown_table(doc, table_lines)
            continue
            
        # 5. Heading parsing
        m_head = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m_head:
            level = len(m_head.group(1))
            text = m_head.group(2).strip()
            
            # Clean up potential markdown formatting in heading text
            clean_text = text.replace("**", "").replace("*", "").replace("`", "")
            
            p = doc.add_heading(level=level)
            parse_and_add_runs(p, clean_text)
            
            # Apply heading specific spacing/page break rules
            p.paragraph_format.keep_with_next = True
            if level == 1:
                # Always start Heading 1 on a new page (unless we are at the very first Heading 1 of body, 
                # but in BCA format, new chapters on new page is perfect)
                p.paragraph_format.page_break_before = True
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
            elif level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
            else:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                
            i += 1
            continue
            
        # 6. Horizontal Rules
        if line.strip() == "---":
            if in_front_matter:
                # Horizontal lines in front matter act as page boundaries
                doc.add_page_break()
            else:
                # Horizontal lines in body act as styled separators
                add_horizontal_rule(doc)
            i += 1
            continue
            
        # 7. Unordered Lists
        m_bullet = re.match(r'^[\*\-]\s+(.*)$', line)
        if m_bullet:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            parse_and_add_runs(p, m_bullet.group(1).strip())
            i += 1
            continue
            
        # 8. Ordered Lists
        m_num = re.match(r'^(\d+)\.\s+(.*)$', line)
        if m_num:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            parse_and_add_runs(p, m_num.group(2).strip())
            i += 1
            continue
            
        # 9. Standalone Images
        m_img = re.match(r'^!\[(.*?)\]\((.*?)\)$', line.strip())
        if m_img:
            alt_text = m_img.group(1).strip()
            img_path = m_img.group(2).strip()
            print(f"Adding image: {img_path}")
            add_image(doc, img_path, alt_text)
            i += 1
            continue
            
        # 10. Standard Paragraphs
        if line.strip():
            p = doc.add_paragraph()
            parse_and_add_runs(p, line.strip())
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
        i += 1

    print(f"Saving compiled docx to: {docx_path}")
    doc.save(docx_path)
    print("DOCX Compilation complete!")

if __name__ == "__main__":
    main()
