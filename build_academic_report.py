import os
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_color):
    """Set the background color of a cell."""
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def add_page_number(run):
    """Add a dynamic page number field to a run in footer."""
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

def add_toc(paragraph):
    """Add a dynamic Table of Contents field to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    placeholder = paragraph.add_run("Right-click this field and select 'Update Field' to generate Table of Contents.")
    placeholder.font.italic = True
    placeholder.font.color.rgb = RGBColor(128, 128, 128)
    
    run2 = paragraph.add_run()
    run2._r.append(fldChar3)

def add_horizontal_rule(doc):
    """Add a thin horizontal divider line."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "CCCCCC")
    
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trHeight = parse_xml(r'<w:trHeight %s w:val="20" w:hRule="exact"/>' % nsdecls('w'))
    trPr.append(trHeight)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def parse_and_add_runs(paragraph, text):
    """Parse basic markdown inline styles (**bold**, *italic*, `code`) and append runs."""
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|_.*?_)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("_") and part.endswith("_"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(197, 17, 98)
        else:
            run = paragraph.add_run(part)

def add_code_block(doc, code_text, filename=""):
    """Render a code block inside a beautiful single-cell table callout container."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8F9FA")
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="120" w:type="dxa"/>'
        f'<w:left w:w="180" w:type="dxa"/>'
        f'<w:bottom w:w="120" w:type="dxa"/>'
        f'<w:right w:w="180" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)
    
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="E4E7EC"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="A4BCDE"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E4E7EC"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="E4E7EC"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(34, 34, 34)
    
    # Spacing after code block
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

def add_styled_table(doc, headers, rows_data):
    """Add a structured, styled table with headers and light zebra striping."""
    num_cols = len(headers)
    num_rows = len(rows_data) + 1
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Style header row
    for col_idx, text in enumerate(headers):
        cell = table.cell(0, col_idx)
        set_cell_background(cell, "1F4E79")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Style data rows
    for row_idx, r_data in enumerate(rows_data):
        actual_row = row_idx + 1
        for col_idx, text in enumerate(r_data):
            cell = table.cell(actual_row, col_idx)
            if actual_row % 2 == 1:
                set_cell_background(cell, "F2F5F8")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            parse_and_add_runs(p, str(text))
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)

def add_image(doc, img_path, alt_text):
    """Load and center an image, adding a structured caption underneath."""
    base_dir = r"a:\Interview_Generator"
    full_path = os.path.join(base_dir, img_path.replace("/", "\\"))
    
    if os.path.exists(full_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(full_path, width=Inches(5.8))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run(f"Figure: {alt_text}")
        run_cap.italic = True
        run_cap.font.size = Pt(9.5)
        run_cap.font.color.rgb = RGBColor(100, 100, 100)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Image Missing: {alt_text} at {img_path}]")
        run.bold = True
        run.font.color.rgb = RGBColor(200, 0, 0)

def extract_codebase(md_path):
    """Parse existing report file and extract code files from section 6."""
    if not os.path.exists(md_path):
        return []
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    code_blocks = []
    current_file = None
    current_lines = []
    in_code = False
    
    for line in lines:
        if line.startswith("### 6."):
            parts = line.split()
            if len(parts) >= 3:
                current_file = parts[2]
                current_file = current_file.split("(")[0].strip()
        elif line.strip().startswith("```") and current_file:
            if not in_code:
                in_code = True
                current_lines = []
            else:
                in_code = False
                code_blocks.append((current_file, "".join(current_lines)))
                current_file = None
        elif in_code:
            current_lines.append(line)
            
    return code_blocks

def configure_styles(doc):
    """Set global document styles for Calibri consistency."""
    styles = doc.styles
    style_normal = styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(34, 34, 34)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    style_h1 = styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Calibri'
    font_h1.size = Pt(18)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(31, 78, 121)
    style_h1.paragraph_format.space_before = Pt(18)
    style_h1.paragraph_format.space_after = Pt(8)
    style_h1.paragraph_format.keep_with_next = True
    
    style_h2 = styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Calibri'
    font_h2.size = Pt(14)
    font_h2.bold = True
    font_h2.color.rgb = RGBColor(47, 85, 151)
    style_h2.paragraph_format.space_before = Pt(14)
    style_h2.paragraph_format.space_after = Pt(6)
    style_h2.paragraph_format.keep_with_next = True
    
    style_h3 = styles['Heading 3']
    font_h3 = style_h3.font
    font_h3.name = 'Calibri'
    font_h3.size = Pt(12)
    font_h3.bold = True
    font_h3.color.rgb = RGBColor(89, 89, 89)
    style_h3.paragraph_format.space_before = Pt(10)
    style_h3.paragraph_format.space_after = Pt(4)
    style_h3.paragraph_format.keep_with_next = True

def add_heading_with_break(doc, text, level=1):
    """Add a heading, ensuring Level 1 starts on a new page."""
    p = doc.add_heading(level=level)
    parse_and_add_runs(p, text)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_paragraph(doc, text):
    """Add a normal paragraph with standard spacing."""
    p = doc.add_paragraph()
    parse_and_add_runs(p, text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p

def main():
    md_path = r"a:\Interview_Generator\PrepEdge_AI_Project_Report.md"
    docx_path = r"a:\Interview_Generator\PrepEdge_AI_Project_Report.docx"
    
    print("Extracting codebase from previous report...")
    codebase = extract_codebase(md_path)
    print(f"Extracted {len(codebase)} source code files.")
    
    print("Creating new document...")
    doc = Document()
    configure_styles(doc)
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.different_first_page_header_footer = True
    
    # Headers & Footers
    header = section.header
    p_header = header.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = p_header.add_run("PrepEdge AI — Academic Project Report")
    run_h.font.name = "Calibri"
    run_h.font.size = Pt(8.5)
    run_h.font.color.rgb = RGBColor(128, 128, 128)
    
    footer = section.footer
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run()
    add_page_number(run_f)
    run_f.font.name = "Calibri"
    run_f.font.size = Pt(9.5)
    run_f.font.color.rgb = RGBColor(128, 128, 128)
    
    # Cover Page
    for _ in range(2):
        doc.add_paragraph()
        
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("A PROJECT REPORT ON")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("PREPEDGE AI: AN AI-POWERED INTERVIEW SIMULATOR &\nMOCK EVALUATION PLATFORM")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run(
        "Submitted in partial fulfillment of the requirements\n"
        "for the degree of\n"
        "Bachelor of Computer Applications (BCA)"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    add_horizontal_rule(doc)
    doc.add_paragraph()
    
    # Cover Metadata
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(3.0)
        
    meta = [
        ("Submitted By:", "[Your Name]\nRoll No: [Your Roll Number]"),
        ("Guided By:", "[Supervisor Name]\nAssistant Professor"),
        ("Department:", "Department of Computer Applications"),
        ("Institution:", "[Your College/Institution Name]\nAcademic Year: 2025-2026")
    ]
    for idx, (label, val) in enumerate(meta):
        c0 = table.cell(idx, 0)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(label)
        r0.font.name = "Calibri"
        r0.font.size = Pt(10.5)
        r0.bold = True
        r0.font.color.rgb = RGBColor(80, 80, 80)
        
        c1 = table.cell(idx, 1)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(val)
        r1.font.name = "Calibri"
        r1.font.size = Pt(10.5)
        
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="none"/>'
                f'<w:left w:val="none"/>'
                f'<w:bottom w:val="none"/>'
                f'<w:right w:val="none"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)
            
    doc.add_page_break()
    
    # Certificate page
    p = doc.add_heading(level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CERTIFICATE OF ORIGINALITY")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    
    add_paragraph(doc, 
        "This is to certify that the project report entitled **\"PrepEdge AI: An AI-Powered Interview Simulator & Mock Evaluation Platform\"** "
        "submitted to the Department of Computer Applications is a record of original work done by the candidate under supervision and guidance."
    )
    add_paragraph(doc, 
        "The work has not been submitted in part or full for any other degree or diploma to any other University or Institution."
    )
    doc.add_paragraph().paragraph_format.space_before = Pt(72)
    
    # Signature placeholders
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.rows[0].cells[0].width = Inches(3.0)
    sig_table.rows[0].cells[1].width = Inches(3.0)
    
    c0 = sig_table.cell(0, 0)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p0.add_run("_____________________\nProject Supervisor")
    r0.font.name = "Calibri"
    r0.font.size = Pt(11)
    
    c1 = sig_table.cell(0, 1)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run("_____________________\nHead of Department")
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    
    for cell in sig_table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/>'
            f'</w:tcBorders>'
        ))
        
    doc.add_page_break()
    
    # Acknowledgements page
    p = doc.add_heading(level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ACKNOWLEDGEMENTS")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    
    add_paragraph(doc, 
        "I express my deepest gratitude to my project supervisor, faculty members, and the department head for their invaluable guidance, "
        "encouragement, and support throughout the design, development, and implementation of this project."
    )
    add_paragraph(doc, 
        "I also extend my sincere thanks to my family, friends, and peers who directly or indirectly assisted in the completion of this software engineering project. "
        "Their timely suggestions and testing feedback were vital in refining the system functionalities."
    )
    doc.add_page_break()
    
    # Abstract page
    p = doc.add_heading(level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ABSTRACT")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    
    add_paragraph(doc, 
        "Modern recruitment processes place a high premium on technical competency, communication clarity, and situational adaptability. "
        "Traditional interview preparation methods lack personalization, interactive simulation, and objective feedback. **PrepEdge AI** is an advanced web-based platform designed to bridge this gap."
    )
    add_paragraph(doc, 
        "By leveraging cutting-edge Artificial Intelligence (Google Gemini and Groq Llama-3 API integrations), PrepEdge AI simulates professional, role-specific mock interviews based either on custom parameters (job title, experience, difficulty) or direct analysis of uploaded resume files."
    )
    add_paragraph(doc, 
        "The system incorporates the browser's Web Speech API for real-time speech-to-text conversion, enabling candidates to answer questions verbally in a timed environment. Each answer is evaluated by the AI model on a percentage scale and categorized as Excellent, Good, Partial, or Needs Work, accompanied by coaching tips and an ideal response representation."
    )
    add_paragraph(doc, 
        "The backend is built with Python and Flask, using MySQL for persistent user sessions, notifications, and activity tracking. Detailed reports are available for export in TXT and PDF formats, and can be emailed to users via SMTP. A fully featured Admin Panel provides user management controls, notifications, and contact resolution portals."
    )
    doc.add_page_break()
    
    # Table of contents page
    p = doc.add_heading(level=1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.page_break_before = False
    p.add_run("TABLE OF CONTENTS")
    add_toc(p)
    doc.add_page_break()
    
    # --- 1. INTRODUCTION ---
    add_heading_with_break(doc, "1. INTRODUCTION", 1)
    add_paragraph(doc, 
        "In the current competitive job market, candidates often struggle with mock interviews due to the lack of live partners and structured evaluation systems. "
        "PrepEdge AI addresses this challenge by providing an autonomous, 24/7 accessible mock interview environment powered by Large Language Models (LLMs)."
    )
    
    add_heading_with_break(doc, "1.1 Objective", 2)
    add_paragraph(doc, 
        "The primary objective of PrepEdge AI is to democratize professional mock interview preparation. It accomplishes this through the following sub-objectives:"
    )
    p_list = [
        "**Custom Scenarios**: Enable users to generate tailored sets of 5-20 questions for any professional job role.",
        "**ATS Resume Processing**: Automatically scan resumes (PDF/DOCX) for critical keywords, score them against target roles, and generate questions matching candidate profiles.",
        "**Speech-Enabled Interface**: Capture user spoken answers live using browser-based Speech Recognition to mimic natural conversation.",
        "**Actionable Evaluation**: Score each answer using LLMs based on key points covered, providing clear feedback, coaching tips, and a sample ideal answer.",
        "**Progress Tracking**: Build dashboard visualizations tracking scores, session history, and weekly engagement statistics.",
        "**Admin Audit Trail**: Develop a secure admin portal to monitor system analytics, audit activity logs, block disruptive users, and resolve contact queries."
    ]
    for item in p_list:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_after = Pt(3)
        parse_and_add_runs(p_item, item)
        
    add_heading_with_break(doc, "1.2 Project Category: Python Web Application", 2)
    add_paragraph(doc, 
        "PrepEdge AI falls under the category of **Artificial Intelligence-Powered Web Applications**. The application uses a Python backend (Flask) for session orchestration and business logic. "
        "The AI component integrates Gemini and Groq APIs to perform advanced generative tasks, including real-time interview profiling, dynamic question generation, ATS resume scanning, and semantic transcript evaluation."
    )
    
    add_heading_with_break(doc, "1.3 Advantages of Python", 2)
    add_paragraph(doc, 
        "Python was selected as the core programming language for this project due to its widespread adoption, extensive ecosystem, and specific advantages in AI engineering:"
    )
    advantages = [
        "**Rich Library Ecosystem**: Access to libraries like PyMySQL, ReportLab, python-docx, and official SDKs for generative AI makes integrations seamless.",
        "**Rapid Prototyping**: Python's clean and expressive syntax allows developers to write complex backend routes and failover routines with minimal boilerplate.",
        "**Excellent AI Integration**: Leading AI providers (Google, Groq, OpenAI) prioritize Python SDKs, ensuring first-class support and access to the latest model capabilities.",
        "**Cross-Platform Deployments**: Python runs identically on Windows, Linux, and macOS environments, facilitating simple development-to-production transitions."
    ]
    for adv in advantages:
        p_adv = doc.add_paragraph(style='List Bullet')
        p_adv.paragraph_format.space_after = Pt(3)
        parse_and_add_runs(p_adv, adv)
        
    # --- 2. TOOLS/PLATFORM REQUIRED ---
    add_heading_with_break(doc, "2. TOOLS/PLATFORM REQUIRED", 1)
    add_paragraph(doc, 
        "Deploying and running PrepEdge AI requires specific hardware and software tools on both the server/developer hosting environment and the client machine."
    )
    
    add_heading_with_break(doc, "2.1 Hardware Requirements", 2)
    add_styled_table(doc, 
        ["Component", "Server Minimum Requirement", "Client Minimum Requirement"],
        [
            ["Processor", "Dual Core 2.0 GHz or higher", "Single Core 1.6 GHz or higher"],
            ["RAM", "4 GB Minimum (8 GB Recommended)", "2 GB Minimum"],
            ["Storage", "100 MB free space (excluding database growth)", "Browser Cache only"],
            ["Input Device", "N/A (SSH console)", "Microphone (compatible with Web Speech API)"]
        ]
    )
    
    add_heading_with_break(doc, "2.2 Software Requirements", 2)
    add_styled_table(doc, 
        ["Software Component", "Specification", "Purpose"],
        [
            ["Operating System", "Windows 10/11, Linux (Ubuntu/Debian), or macOS", "Host environments"],
            ["Backend Language", "Python 3.10 or 3.11", "System business logic interpreter"],
            ["Web Framework", "Flask 3.0.0+", "HTTP request routing and session handler"],
            ["Database Server", "MySQL 8.0+ or MariaDB 10.5+", "Persistent user, session, and audit storage"],
            ["Speech Engine", "Browser-native Web Speech API", "Real-time client-side Speech-to-Text translation"],
            ["AI APIs", "Google GenAI API, Groq LLaMA API", "Generative questions and transcript evaluations"],
            ["PDF Compiler", "ReportLab 4.0.7+", "Compiles session performance results to printable formats"]
        ]
    )
    
    # --- 3. PROBLEM DEFINITION ---
    add_heading_with_break(doc, "3. PROBLEM DEFINITION", 1)
    
    add_heading_with_break(doc, "3.1 Problem Statement", 2)
    add_paragraph(doc, 
        "Traditional interview preparation methods are fundamentally limited. Human coaches are expensive and inaccessible to many candidates, "
        "while static practice lists lack real-time interaction and interactive feedback. Existing platforms offer pre-recorded questions or text-only inputs, "
        "failing to evaluate candidate speech pattern authenticity, and lack customized keyword checks against specific resume profiles."
    )
    add_paragraph(doc, 
        "Additionally, candidates suffer from ATS (Applicant Tracking System) screening failures, submitting resumes that do not align with target job roles. "
        "PrepEdge AI solves these problems by creating a secure, 24/7 automated platform integrating ATS checks with timed oral simulator interviews and detailed semantic performance scoring."
    )
    
    add_heading_with_break(doc, "3.2 Requirement Specifications", 2)
    add_paragraph(doc, 
        "**Functional Requirements:**"
    )
    frs = [
        "**User Registration and Secure Login**: Authentication using PBKDF2 hashing, accompanied by SMTP OTP checks to verify registration emails.",
        "**Resume Parsing and ATS Grading**: Text extraction from PDF/DOCX resumes, graded against 30+ roles using keyword density scoring.",
        "**Timed Interview Simulation**: Generation of structured job-specific questions, featuring browser-based speech capture.",
        "**Semantic Scoring**: Evaluation of speech transcripts by Gemini/Groq LLMs, providing score breakdowns, coaching tips, and sample ideal answers.",
        "**Email Dispatch**: Automatic delivery of compiled PDF performance certificates directly to user emails.",
        "**Admin Management console**: Capabilities for admins to view system logs, inspect registrations, and block/unblock disruptive accounts."
    ]
    for fr in frs:
        p_fr = doc.add_paragraph(style='List Bullet')
        p_fr.paragraph_format.space_after = Pt(3)
        parse_and_add_runs(p_fr, fr)
        
    # --- 4. PROJECT PLANNING & SCHEDULING ---
    add_heading_with_break(doc, "4. PROJECT PLANNING & SCHEDULING", 1)
    
    add_heading_with_break(doc, "4.1 Gantt Chart", 2)
    add_paragraph(doc, 
        "The project development lifespan was organized over an **8-week timeline** utilizing Agile cycles. The table below represents the project's Gantt Chart, "
        "outlining the duration and scheduling of each software phase:"
    )
    
    add_styled_table(doc, 
        ["Phase", "Description", "Start Week", "End Week", "Status"],
        [
            ["Requirement Gathering", "SRS formulation and initial user story mappings", "Week 1", "Week 2", "Completed"],
            ["Database & System Design", "ER diagrams, schema setups, and MVC routing design", "Week 2", "Week 3", "Completed"],
            ["Core Backend Coding", "Flask setup, database wrappers, SMTP setups", "Week 3", "Week 5", "Completed"],
            ["AI & Speech Integration", "Gemini/Groq SDK wiring and Web Speech API hooks", "Week 4", "Week 6", "Completed"],
            ["Frontend Design & CSS", "Dashboard layout, responsive screens, SVG widgets", "Week 5", "Week 7", "Completed"],
            ["System Testing & QA", "Unit checks, failover validation, and integration tests", "Week 7", "Week 8", "Completed"],
            ["Deployment & Review", "MySQL optimization, production setups, final documentation", "Week 8", "Week 8", "Completed"]
        ]
    )
    
    add_heading_with_break(doc, "4.2 PERT Chart", 2)
    add_paragraph(doc, 
        "The PERT (Program Evaluation and Review Technique) analysis details the task sequence and dependencies. The critical path runs through "
        "**Requirement Gathering -> System Design -> Core Backend -> AI Integration -> System Testing -> Deployment**."
    )
    
    add_styled_table(doc, 
        ["Task ID", "Activity Description", "Predecessors", "Optimistic (O)", "Pessimistic (P)", "Most Likely (M)", "Expected Time (Te)"],
        [
            ["A", "Requirement Gathering", "None", "5 days", "15 days", "10 days", "10.0 days"],
            ["B", "System Design & ERD", "A", "4 days", "10 days", "6 days", "6.3 days"],
            ["C", "Core Backend Development", "B", "10 days", "22 days", "15 days", "15.3 days"],
            ["D", "Speech API & AI Integration", "C", "7 days", "18 days", "12 days", "12.2 days"],
            ["E", "Frontend Layout & Styling", "C", "6 days", "14 days", "9 days", "9.3 days"],
            ["F", "Testing and QA Verification", "D, E", "5 days", "12 days", "8 days", "8.1 days"],
            ["G", "Deployment and Presentation", "F", "2 days", "6 days", "4 days", "4.0 days"]
        ]
    )
    add_paragraph(doc, "*Expected Time formula used: Te = (O + 4M + P) / 6*")
    
    # --- 5. SYSTEM ANALYSIS ---
    add_heading_with_break(doc, "5. SYSTEM ANALYSIS", 1)
    
    add_heading_with_break(doc, "5.1 Identification of Need", 2)
    add_paragraph(doc, 
        "Modern HR screening practices increasingly utilize automated phone screenings and initial ATS keyword check loops. "
        "Candidates fail because they cannot practice in an environment that simulates these automated constraints. "
        "The identification of need focuses on providing a **safe, feedback-oriented simulator** that:"
    )
    needs = [
        "Reduces test anxiety by normalizing timed speech responses.",
        "Aligns student qualifications with actual corporate expectations prior to job applications.",
        "Offers quantitative indicators (percentages) and qualitative coaching advice to guide continuous skill refinement."
    ]
    for need in needs:
        p_need = doc.add_paragraph(style='List Bullet')
        p_need.paragraph_format.space_after = Pt(3)
        parse_and_add_runs(p_need, need)
        
    add_heading_with_break(doc, "5.2 Feasibility Study", 2)
    add_paragraph(doc, 
        "A multi-dimensional feasibility study was conducted to ensure the project's viability."
    )
    
    add_heading_with_break(doc, "5.2.1 Technical Feasibility", 3)
    add_paragraph(doc, 
        "The project is highly feasible technically. The browser-native **Web Speech API** removes client-side installation requirements (no external binary codecs needed). "
        "Integrating Google GenAI and Groq SDKs within Flask is straightforward. The failover logic automatically switches to LLaMA-3 over Groq if Gemini hits rate limits, "
        "ensuring system resilience."
    )
    
    add_heading_with_break(doc, "5.2.2 Economical Feasibility", 3)
    add_paragraph(doc, 
        "The economic costs are minimal. Python, Flask, and MySQL are free, open-source technologies. Generative AI development models offer generous "
        "free-tier or low-cost tokens. The high value-add (interview preparation, resume grading) combined with zero licensing overhead makes the project economically optimal."
    )
    
    add_heading_with_break(doc, "5.2.3 Operational Feasibility", 3)
    add_paragraph(doc, 
        "PrepEdge AI is operationally intuitive. For candidates, it requires simple clicks to configure interviews and start recording. For admins, "
        "a visual dashboard streamlines account status and notification controls, meaning no advanced training is required to run the system."
    )
    
    # --- 6. SYSTEM ANALYSIS / ANALYSIS ---
    add_heading_with_break(doc, "6. ANALYSIS", 1)
    
    add_heading_with_break(doc, "6.1 Data Flow Diagram (DFD)", 2)
    add_paragraph(doc, 
        "The Data Flow Diagram describes how information moves from entities through system processes into databases and external API systems."
    )
    add_paragraph(doc, 
        "**Context Level DFD (Level 0):**"
    )
    add_styled_table(doc, 
        ["Entity", "Input to System", "Output from System"],
        [
            ["Candidate User", "Registration data, Resume files, Speech Answers, Session configs", "Dashboard stats, Test results, PDF certificate emails"],
            ["Administrator", "Block commands, contact resolutions, credential updates", "Global metrics, audit notification logs, query views"],
            ["AI API Server", "Generated prompts, user transcripts, reference hints", "JSON arrays of questions, scored evaluations"]
        ]
    )
    
    # --- 7. MODULE DESCRIPTION ---
    add_heading_with_break(doc, "7. MODULE DESCRIPTION", 1)
    add_paragraph(doc, 
        "PrepEdge AI consists of five primary integrated software modules:"
    )
    modules = [
        "**1. Authentication & OTP Verification**: Controls registration and forgot-password resets. Securely sends 6-digit OTP codes via SMTP and encrypts user passwords using PBKDF2 hashing.",
        "**2. AI Question Generator & Failover**: Builds context prompts for job categories. First attempts generation via Google Gemini; in case of timeout or quota limits, it falls back seamlessly to LLaMA-3 via Groq.",
        "**3. ATS Resume Scorer & Analyzer**: Extracts text from PDF/DOCX resumes, scans for matching role keywords, calculates keyword density match percentage, and suggests adjustments.",
        "**4. Speech Evaluation Engine**: Integrates Web Speech API inside candidate screens, capturing microphone inputs, sending transcripts to generative AI, and rendering score metrics.",
        "**5. Administrator Control Console**: Provides analytics charts, audit logs, notification feeds, user blocking actions, and message query resolution tools."
    ]
    for mod in modules:
        p_mod = doc.add_paragraph()
        p_mod.paragraph_format.space_after = Pt(4)
        parse_and_add_runs(p_mod, mod)
        
    # --- 8. DATA STRUCTURE SNAPSHOT ---
    add_heading_with_break(doc, "8. DATA STRUCTURE SNAPSHOT", 1)
    add_paragraph(doc, 
        "The backend utilizes a relational database schema designed for speed, security, and reference safety. Below are the table structure snapshots:"
    )
    
    add_heading_with_break(doc, "8.1 Table: users", 2)
    add_styled_table(doc, 
        ["Column Name", "Data Type", "Constraints", "Description"],
        [
            ["id", "INT", "PRIMARY KEY, AUTO_INCREMENT", "Unique identifier for each user"],
            ["username", "VARCHAR(50)", "UNIQUE, NOT NULL", "Login handle for the candidate"],
            ["email", "VARCHAR(100)", "UNIQUE, NOT NULL", "Registered email address"],
            ["password", "VARCHAR(255)", "NOT NULL", "PBKDF2 salted password hash"],
            ["avatar", "VARCHAR(100)", "DEFAULT 'default.png'", "User profile image name"],
            ["blocked", "TINYINT(1)", "DEFAULT 0", "Account status flag (1=blocked)"],
            ["created_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Registration timestamp"]
        ]
    )
    
    add_heading_with_break(doc, "8.2 Table: sessions", 2)
    add_styled_table(doc, 
        ["Column Name", "Data Type", "Constraints", "Description"],
        [
            ["id", "INT", "PRIMARY KEY, AUTO_INCREMENT", "Unique session identifier"],
            ["user_id", "INT", "FOREIGN KEY -> users(id)", "User who initiated the session"],
            ["job_title", "VARCHAR(100)", "NOT NULL", "Target role of the mock interview"],
            ["difficulty", "VARCHAR(20)", "NOT NULL", "Level selected (Easy, Medium, Hard)"],
            ["question_count", "INT", "NOT NULL", "Total number of generated questions"],
            ["created_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Session initiation timestamp"]
        ]
    )
    
    add_heading_with_break(doc, "8.3 Table: question_scores", 2)
    add_styled_table(doc, 
        ["Column Name", "Data Type", "Constraints", "Description"],
        [
            ["id", "INT", "PRIMARY KEY, AUTO_INCREMENT", "Unique answer score ID"],
            ["session_id", "INT", "FOREIGN KEY -> sessions(id)", "Associated interview session"],
            ["question_text", "TEXT", "NOT NULL", "The question generated by the AI"],
            ["hint", "TEXT", "NOT NULL", "Key concepts that should be covered"],
            ["answer_text", "TEXT", "NOT NULL", "The transcribed response text"],
            ["score_label", "VARCHAR(20)", "NOT NULL", "Rating (Excellent, Good, Partial, Needs Work)"],
            ["score_percent", "INT", "NOT NULL", "Percentage grade (0-100)"],
            ["feedback", "TEXT", "NOT NULL", "Detailed review generated by LLM"],
            ["tip", "TEXT", "NOT NULL", "Suggested coaching tip for next time"],
            ["ideal_answer", "TEXT", "NOT NULL", "AI-modeled ideal response reference"]
        ]
    )
    
    # --- 9. LIST OF REPORTS ---
    add_heading_with_break(doc, "9. LIST OF REPORTS", 1)
    add_paragraph(doc, 
        "The system auto-compiles and delivers three high-value reports:"
    )
    reports = [
        "**1. Candidate Performance Certificate (PDF)**: Compiled via ReportLab, this report summarizes the mock interview results, presenting a table of scores, feedback, tips, and an overall percentage grade.",
        "**2. ATS Resume Analysis Summary**: Provides candidates with role keywords matches, density percentages, missing concepts, and optimization score cards.",
        "**3. Administrative Activity Audit Feed**: Generates system alerts regarding new registrations, blocked accounts, API failovers, and resolved contact requests."
    ]
    for rep in reports:
        p_rep = doc.add_paragraph(style='List Bullet')
        p_rep.paragraph_format.space_after = Pt(3)
        parse_and_add_runs(p_rep, rep)
        
    # --- 10. SCOPE AND FUTURE ENHANCEMENT ---
    add_heading_with_break(doc, "10. SCOPE AND FUTURE ENHANCEMENT", 1)
    
    add_heading_with_break(doc, "10.1 Implementation Methodology", 2)
    add_paragraph(doc, 
        "The development team followed an **Agile-Scrum software methodology**. Development occurred in two-week cycles (sprints), starting with sprint planning to select backlogs, "
        "followed by coding, daily standups to identify blockers, and sprint review demonstrations. Continuous integration was utilized to verify database wrappers "
        "and integration states."
    )
    
    add_heading_with_break(doc, "10.2 Future Scope", 2)
    add_paragraph(doc, 
        "Planned extensions for PrepEdge AI include:"
    )
    f_scopes = [
        "**Real-time video/emotion checking**: Integrating camera streams and computer vision checking to analyze posture, eye contact, and expressions.",
        "**Custom Company Templates**: Building mock interview engines aligned with specific recruitment questions of top-tier employers (e.g. Amazon STAR model, Google cognitive checks).",
        "**Local LLM Integration**: Running offline open-source models (like LLaMA-3-8B locally via Ollama) to guarantee absolute data privacy for enterprise clients."
    ]
    for fs in f_scopes:
        p_fs = doc.add_paragraph(style='List Bullet')
        p_fs.paragraph_format.space_after = Pt(3)
        parse_and_add_runs(p_fs, fs)
        
    # --- 11. TESTING ---
    add_heading_with_break(doc, "11. TESTING", 1)
    add_paragraph(doc, 
        "PrepEdge AI was validated against rigorous test specifications to verify system integrations, database consistency, and API recoveries."
    )
    add_styled_table(doc, 
        ["Test Case ID", "Module", "Scenario", "Expected Outcome", "Status"],
        [
            ["TC-01", "Auth", "Register with invalid domain email", "Blocked by domain validation logic", "Passed"],
            ["TC-02", "Auth", "Account block propagation", "Access denied on next request with warning", "Passed"],
            ["TC-03", "Generator", "Job Title validation", "Keyboard mash (e.g., 'asdfgh') is rejected", "Passed"],
            ["TC-04", "Generator", "API Failover", "Gemini failure initiates Groq fallback check", "Passed"],
            ["TC-05", "Resume", "ATS Scorer", "Quantification detection checks for percentage indicators", "Passed"],
            ["TC-06", "Evaluator", "Audio Silence", "Prompt asks user to verify mic permission", "Passed"],
            ["TC-07", "Export", "PDF Compilation", "ReportLab compiles multi-column grid accurately", "Passed"]
        ]
    )
    
    # --- 12. LIMITATIONS AND FURTHER ENHANCEMENTS ---
    add_heading_with_break(doc, "12. LIMITATIONS AND FURTHER ENHANCEMENTS", 1)
    add_paragraph(doc, 
        "While highly functional, the system has minor limitations:"
    )
    limits = [
        "**Internet Dependency**: Generative evaluations and resume ATS scoring require active access to Gemini/Groq APIs, making it offline-incompatible.",
        "**Accent Pronunciation**: Web Speech API speech-to-text accuracy can degrade in the presence of strong accents or background acoustic noise.",
        "**Absence of Non-Verbal Cues**: The system currently scores content transcripts only, missing body language and vocal tone checks."
    ]
    for lim in limits:
        p_lim = doc.add_paragraph(style='List Bullet')
        p_lim.paragraph_format.space_after = Pt(3)
        parse_and_add_runs(p_lim, lim)
        
    # --- 13. SECURITY MECHANISM ---
    add_heading_with_break(doc, "13. SECURITY MECHANISM", 1)
    add_paragraph(doc, 
        "A multi-tier security framework is implemented to guarantee data privacy and protect host operations:"
    )
    securities = [
        "**Password Hashing**: User authentication relies on secure `pbkdf2:sha256` salted hashes, ensuring stored values cannot be decoded.",
        "**Access Control Decorators**: Route protection restricts app functionalities to authenticated user sessions, redirecting unauthenticated requests.",
        "**SQL Injection Guard**: Custom database adapters in `database.py` use parametrized SQL statements to prevent SQL injections.",
        "**OTP Validation**: Registration and forgot-password routines require dynamic OTP verification to prevent spam and verify ownership."
    ]
    for sec in securities:
        p_sec = doc.add_paragraph(style='List Bullet')
        p_sec.paragraph_format.space_after = Pt(3)
        parse_and_add_runs(p_sec, sec)
        
    # --- 14. SNAPSHOTS ---
    add_heading_with_break(doc, "14. SNAPSHOTS", 1)
    add_paragraph(doc, 
        "Below are screenshots showing the responsive user interfaces of the PrepEdge AI application:"
    )
    
    snapshots_data = [
        ["report_images/landing_page.png", "Landing Page - Highlights platform features, statistical widgets, and testimonials."],
        ["report_images/login_page.png", "Login Portal - Role-based tabs allowing both users and administrators to authenticate securely."],
        ["report_images/register_page.png", "OTP Registration Form - Email checks and SMTP OTP validation screen."],
        ["report_images/main_app_page.png", "Main Workspace - Controls to configure interviews, upload resumes, and review parameters."],
        ["report_images/mock_interview.png", "Mock Interview Session - Speech capture screen with timer, mic controls, and AI prompt transcripts."],
        ["report_images/admin_dashboard.png", "Admin Panel - Dashboard tracking global usage statistics, user statuses, and alert logs."]
    ]
    for img_path, desc in snapshots_data:
        add_image(doc, img_path, desc)
        
    # --- 15. BIBLIOGRAPHY & REFERENCES ---
    add_heading_with_break(doc, "15. BIBLIOGRAPHY & REFERENCES", 1)
    
    add_heading_with_break(doc, "15.1 Bibliography", 2)
    add_paragraph(doc, "1. Grinberg, M. (2018). *Flask Web Development: Developing Web Applications with Python*. O'Reilly Media.")
    add_paragraph(doc, "2. McKinney, W. (2012). *Python for Data Analysis*. O'Reilly Media.")
    add_paragraph(doc, "3. Elmasri, R., & Navathe, S. B. (2015). *Fundamentals of Database Systems*. Pearson.")
    
    add_heading_with_break(doc, "15.2 References", 2)
    add_paragraph(doc, "1. Flask API Documentation: https://flask.palletsprojects.com/")
    add_paragraph(doc, "2. Google Generative AI Python SDK: https://github.com/google/generative-ai-python")
    add_paragraph(doc, "3. Groq SDK Python Reference Guide: https://github.com/groq/groq-python")
    add_paragraph(doc, "4. MDN Web Docs: Web Speech API Spec: https://developer.mozilla.org/en-US/docs/Web/API/Web_Speech_API")
    add_paragraph(doc, "5. ReportLab PDF Generation Library Guide: https://www.reportlab.com/documentation/")
    
    # --- 16. APPENDIX: FULL IMPLEMENTATION CODEBASE ---
    add_heading_with_break(doc, "16. APPENDIX: FULL IMPLEMENTATION CODEBASE", 1)
    add_paragraph(doc, 
        "This section documents the complete production source files used to build, secure, and deploy the PrepEdge AI system. "
        "The codebase exhibits strict modularity and defensive coding standards:"
    )
    
    for filename, code_content in codebase:
        add_heading_with_break(doc, f"Source File: {filename}", 2)
        add_code_block(doc, code_content, filename)
        
    print(f"Saving restructured academic docx to: {docx_path}")
    doc.save(docx_path)
    print("Academic DOCX successfully compiled!")

if __name__ == "__main__":
    main()
